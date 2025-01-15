from docx import Document
from docx.shared import Mm
from lxml import etree
from bs4 import BeautifulSoup
from io import StringIO

class TABLE2DOC():
    def __init__(self, doc, table_html_str, out_put_path=None):
        self.out_put_path = out_put_path  # 适合单独表格
        self.doc = doc if doc else Document()
        self.table, self.merge_array = self.ana_table_size(table_html_str)

    def save_table(self):
        if not self.out_put_path:
            return 0
            # 保存doc
        self.doc.save(self.out_put_path)

    def get_inch(self):
        # 根据输入内容计算单元格尺寸大小
        price = 0
        return Mm(price)

    def ana_table_size(self, table_html_str):
        # 根据输入内容解析表格最大行列值。
        tree = etree.parse(StringIO(table_html_str), parser=etree.HTMLParser())
        # 获取所有表格行
        rows = tree.xpath('//table/tr')
        # 遍历表格行，判断最大列，累加最大行
        col_count_list = []
        merge_array = []
        row_count = len(rows)
        for i, row in enumerate(rows):
            colspan = 0
            # 获取每一行中的所有单元格
            cells = row.xpath('td')
            # 遍历单元格，判断最大列
            for j, cell in enumerate(cells):

                colspan_count = int(cell.get('colspan', 1))
                rowspan_count = int(cell.get('rowspan', 1))
                colspan += colspan_count
                for index in range(1, colspan_count):
                    merge_array.append([[i, j], [i, j + index]])
                for index in range(1, rowspan_count):
                    merge_array.append([[i, j], [i + index, j]])
            col_count_list.append(colspan)
        col_count = max(col_count_list)
        print(row_count, col_count)
        table = self.doc.add_table(rows=row_count, cols=col_count, style="Table Grid")
        table.autofit = False
        # print(merge_array)
        return table, merge_array

    def parse_table2list(self, table_html_str):

        soup = BeautifulSoup(table_html_str, 'html.parser')
        table = soup.find('table')
        rows = table.find_all('tr')
        # 获取表格的最大行数和列数
        max_cols = 0
        for row in rows:
            cols = row.find_all(['td', 'th'])
            max_cols = max(max_cols, len(cols))
        # 初始化二维列表
        data = [[None] * max_cols for _ in range(len(rows))]
        # 填充二维列表
        row_index = 0
        for row in rows:
            col_index = 0
            cols = row.find_all(['td', 'th'])
            for col in cols:
                rowspan = int(col.get('rowspan', 1))
                colspan = int(col.get('colspan', 1))
                # 找到第一个空位置
                while col_index < max_cols and data[row_index][col_index] is not None:
                    col_index += 1
                # 填充数据
                for r in range(row_index, row_index + rowspan):
                    for c in range(col_index, col_index + colspan):
                        if r < len(data) and c < max_cols:
                            data[r][c] = col.get_text(strip=True)
                # 更新列索引
                col_index += colspan
            # 更新行索引
            row_index += 1
        return data

    # 判断是否合并单元格
    def merge_cell(self, cell1_list, cell2_list):
        print(cell1_list, cell2_list)
        try:
            start_cell = self.table.cell(cell1_list[0], cell1_list[1])
            end_cell = self.table.cell(cell2_list[0], cell2_list[1])
            end_cell.text = ""
            start_cell.merge(end_cell)
        except:
            pass

    def main_generate_table(self, table_html_str):
        merge_array = self.merge_array

        data = self.parse_table2list(table_html_str)
        # 表格赋值
        for i, rows in enumerate(self.table.rows):  # 获取所有行
            for j, cell in enumerate(rows.cells):  # 根据行遍历列
                cell.text = data[i][j]
        # 记录已经合并过的单元格
        previous_end = None

        # 合并单元格
        for i in range(len(merge_array)):
            start, end = merge_array[i]
            if previous_end is not None and start == previous_end:
                # 如果开始索引与上一次的末尾索引值相同，对后续的所有值进行加1处理
                for j in range(i, len(merge_array)):
                    merge_array[j][0][1] += 1
                    merge_array[j][1][1] += 1
                    previous_end = merge_array[j][1]
            else:
                previous_end = end
        print(merge_array)
        for merge in merge_array:
            cell1_list, cell2_list = merge
            self.merge_cell(cell1_list, cell2_list)
        return self.table


if __name__ == '__main__':
    table_html_str = '<html><body><table><tr><td colspan="6">网站系统安全技术措施</td></tr><tr><td colspan="2">层面防护</td><td colspan="4">整体防护</td></tr><tr><td>网站层</td><td>Web应用安全 城名安全</td><td rowspan="6">运行支撑</td><td rowspan="6">攻击防范</td><td rowspan="6">安全监控</td><td rowspan="6">应急响应</td></tr><tr><td>数据层</td><td>内容发布及数据安全</td></tr><tr><td>主机层</td><td>服务器安全</td></tr><tr><td>网络层</td><td>管理终端安全 边界安全</td></tr><tr><td></td><td></td></tr><tr><td>物理层</td><td>物理安全</td></tr></table></body></html>'
    # [[[0, 0], [0, 1]], [[0, 0], [0, 2]], [[0, 0], [0, 3]], [[0, 0], [0, 4]], [[0, 0], [0, 5]], [[1, 0], [1, 1]], [[1, 1], [1, 2]], [[1, 1], [1, 3]], [[1, 1], [1, 4]], [[2, 2], [3, 2]], [[2, 2], [4, 2]], [[2, 2], [5, 2]], [[2, 2], [6, 2]], [[2, 2], [7, 2]], [[2, 3], [3, 3]], [[2, 3], [4, 3]], [[2, 3], [5, 3]], [[2, 3], [6, 3]], [[2, 3], [7, 3]], [[2, 4], [3, 4]], [[2, 4], [4, 4]], [[2, 4], [5, 4]], [[2, 4], [6, 4]], [[2, 4], [7, 4]], [[2, 5], [3, 5]], [[2, 5], [4, 5]], [[2, 5], [5, 5]], [[2, 5], [6, 5]], [[2, 5], [7, 5]]]
    doc = Document()
    t2d = TABLE2DOC(doc, table_html_str, 'test.docx')
    t2d.main_generate_table(table_html_str)
    t2d.save_table()
