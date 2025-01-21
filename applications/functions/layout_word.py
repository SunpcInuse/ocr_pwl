import os

import cv2
from pathlib import Path
from rapid_layout import RapidLayout, VisLayout
from rapidocr_onnxruntime import RapidOCR
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rapid_table import RapidTable, VisTable

table_engine = RapidTable()
ocr_engine = RapidOCR()

from applications.functions.seg_table import TABLE2DOC


class LayoutWord():
    def __init__(self, save_path=None, conf_threshold=0.7):
        self.layout_engine = RapidLayout(model_type="pp_layout_cdla", conf_thres=conf_threshold)
        self.engine = RapidOCR()
        self.general_category_list = ['title', 'text', 'figure', 'figure_caption', 'table', 'table_caption', 'header',
                                      'footer',
                                      'reference', 'equation']
        self.vis_save_path = ""
        self.inference_path = Path(r'inference_results')
        self.inference_path.mkdir(parents=True, exist_ok=True)
        self.save_path = save_path

    def layout_pic(self, image_paths: list):
        # if isinstance(image_path, str):
        #     image = cv2.imread(image_path)
        # else:
        #     image = image_path
        image_list, ret_res = [], []
        for image_path in image_paths:
            image = cv2.imread(image_path)
            image_list.append(image)
            original_height, original_width, _ = image.shape
            # print(original_height, original_width)
            boxes, scores, class_names, elapse = self.layout_engine(image)
            # ploted_img = VisLayout.draw_detections(image, boxes, scores, class_names)
            # save_layout_path = self.inference_path / "layout_res{}.png".format(self.vis_save_path)
            # print(save_layout_path)
            # if ploted_img is not None:  # 将版面分析的结果可视化
            #     cv2.imwrite(save_layout_path, image)
            detection_results = []
            for box, score, class_name in zip(boxes, scores, class_names):
                x1, y1, x2, y2 = box  # x1: 左上角x坐标, y1: 左上角y坐标, x2: 右下角x坐标, y2: 右下角y坐标
                x1, y1, x2, y2 = int(x1), int(y1), int(x2), int(y2)
                cropped_image = image[y1: y2, x1: x2]
                content = self.extract_text(cropped_image)
                detection_results.append([class_name, content, [x1, y1, x2, y2]])
            # print(detection_results)
            detection_results.sort(key=lambda x: x[2][1])
            ret_res.append(detection_results)
        return image_list, ret_res

    def extract_text(self, cropped_image):
        result, _ = self.engine(cropped_image)
        content = ""
        if result is None:
            return content
        for _res in result:
            content += _res[1]  # 正确拼接字符串
        return content

    def header_footer_extract(self, image, detection_results):
        cropped_images = []
        header_content = ""
        footer_content = ""
        for detection in detection_results:
            box, score, class_name = detection[0], detection[1], detection[2]
            # print(box, score, class_name)
            x1, y1, x2, y2 = box
            x1, y1, x2, y2 = int(x1), int(y1), int(x2), int(y2)
            cropped_image = image[y1: y2, x1: x2]
            cropped_images.append(cropped_image)

            if class_name == "header":
                header_content = self.extract_text(cropped_image)
            elif class_name == "footer":
                footer_content = self.extract_text(cropped_image)

        return header_content, footer_content

    def figure_table_extract(self, image):
        pass

    def reference_equation_extract(self, image):
        pass

    def layout_crop_detection(self, image, detection_results):
        cropped_images = []
        header_content = ""
        footer_content = ""
        for detection in detection_results:
            box, score, class_name = detection[0], detection[1], detection[2]
            # print(box, score, class_name)
            x1, y1, x2, y2 = box
            x1, y1, x2, y2 = int(x1), int(y1), int(x2), int(y2)
            cropped_image = image[y1: y2, x1: x2]
            cropped_images.append(cropped_image)

            if class_name == "header":
                header_content = self.extract_text(cropped_image)
            elif class_name == "footer":
                footer_content = self.extract_text(cropped_image)

    def determine_title_level(self, title):
        if not title:
            return 0
        level = title.strip().count('.') + 1
        return level

    def format_info(self, layout_res: list):
        if layout_res is None:
            return None
        # 判断版面分析的结果分类，针对不同类别制作对应的版面信息。
        # 类编包括 title、text、 figure figure_caption table table_caption header footer reference equation
        if layout_res[0] in ['title', 'text', 'reference', 'figure_caption', 'footer', 'table_caption', 'header']:
            bbox = layout_res[2]
            text = layout_res[1]
            # 根据content数量和box的上下边界，判断字体大小
            """
                根据文本框的高度和宽度推算字体大小。
            """
            # print(bbox)
            text_width = bbox[2] - bbox[0]  # 文本框的宽度
            text_height = bbox[3] - bbox[1]  # 文本框的高度

            font_size_based_on_height = text_height  # 假设字体的高度与文本框高度一致

            # 平均字符宽度，假设字符宽度为文本框宽度除以字符数
            avg_char_width = text_width / len(text) if len(text) > 0 else text_width  # 每个字符的平均宽度
            font_size_based_on_width = text_width * 0.9 / (len(text) * avg_char_width) * font_size_based_on_height

            font_size = min(font_size_based_on_height, font_size_based_on_width)
            return font_size * 0.075
        else:
            return None

    def generate_word_document(self, image_paths: list):
        # print(image_paths)
        # 创建一个新的 Word 文档
        doc = Document()
        # 遍历每个图像路径
        image_list, ret_res = self.layout_pic(image_paths)
        # 遍历检测结果
        for image, total_detection in zip(image_list, ret_res):
            # 获取图像和检测结果
            for detection in total_detection:
                class_name, content, box = detection
                if class_name == 'text':
                    # 添加文本
                    text = doc.add_paragraph()
                    # font_size = self.format_info(detection)
                    r = text.add_run(content)
                    r.font.size = Pt(10.5)
                elif class_name == 'title':
                    # 添加标题
                    level = self.determine_title_level(content)
                    if level > 0:
                        doc.add_heading(content, level=level)
                    else:
                        doc.add_heading(content, level=5)
                elif class_name == 'figure':
                    # 添加图像
                    x1, y1, x2, y2 = box
                    cropped_image = image[y1:y2, x1:x2]
                    cv2.imwrite("temp_figure.png", cropped_image)
                    doc.add_picture("temp_figure.png", width=Inches(4))
                elif class_name == 'figure_caption':
                    # 添加图像标题
                    figure_caption = doc.add_paragraph()
                    r = figure_caption.add_run(content)
                    r.font.size = Pt(10.5)
                    figure_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif class_name == 'table':
                    # 添加表格
                    x1, y1, x2, y2 = box
                    cropped_image = image[y1:y2, x1:x2]
                    cv2.imwrite("temp_table.png", cropped_image)
                    ocr_result, _ = ocr_engine("temp_table.png")
                    table_html_str, table_cell_bboxes, _ = table_engine("temp_table.png", ocr_result)
                    # print(table_html_str)
                    tbl_engine = TABLE2DOC(doc, table_html_str)
                    tbl_engine.main_generate_table(table_html_str)
                elif class_name == 'table_caption':
                    # 添加表格标题，使用自定义样式居中对齐
                    table_caption_paragraph = doc.add_paragraph()
                    r = table_caption_paragraph.add_run(content)
                    r.font.size = Pt(10.5)
                    table_caption_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif class_name == 'header':
                    # 添加页眉
                    section = doc.sections[0]
                    header = section.header
                    header_paragraph = header.paragraphs[0]
                    header_paragraph.text = content
                elif class_name == 'footer':
                    # 添加页脚
                    section = doc.sections[0]
                    footer = section.footer
                    footer_paragraph = footer.paragraphs[0]
                    footer_paragraph.text = content
                elif class_name == 'reference':
                    # 添加参考文献
                    doc.add_paragraph(content, style='ListNumber')
                elif class_name == 'equation':
                    # 添加公式
                    doc.add_paragraph(content, style='Quote')
        if self.save_path:
            # 保存 Word 文档
            print("Word 文档已保存为:", self.save_path, '.docx')
            doc.save(r'D:\Project\OCR-demo\downloads/{}.docx'.format(self.save_path))
        else:
            doc.save("doc.docx")


if __name__ == '__main__':
    layout_word = LayoutWord("gen")
    layout_word.generate_word_document(["b.jpg", "b-1.jpg"])
